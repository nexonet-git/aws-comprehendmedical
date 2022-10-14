import anvil.email
import anvil.tables as tables
import anvil.tables.query as q
from anvil.tables import app_tables
import anvil.server
import anvil.media
import anvil.pdf
import boto3
import numpy as np
import pandas as pd

import csv
import itertools

import datetime
import docx
import re
import fitz


dictMedicalCondition = { "condition":[]}
dictMedicalProcedure = { "procedure":[]}
dictMedicationGrid = { "medication":[]}
#lstCondition1 = []
lstCondition = []
lstProcedure = []
lstMedication = []
strUserEmail = ''
blnLastDictItemProcedure = 0
Conditions_Data = { "condition":[]}
Medications_Data = { "procedure":[]}
Procedures_Data = { "medication":[]}


@anvil.server.callable
def strip_Text(strMedText):
  #Clean up the medical text, get rid of carriage returns etc.
  txtStripped = re.sub('[^A-Za-z0-9\.\,\:\;\-\(\)\\/ ]+', '#', strMedText)
  txtStripped = txtStripped.replace('#', '. ', )
  txtStripped = txtStripped.replace('..', '.', )
  txtStripped = txtStripped.replace('. .', '.', )

  #Split the text by full stops and store each sentence in a list
  listedText = txtStripped.split(".")

  #Since Comprehend Medical can't distinguish between the patient's medical info and family member's medical info,
  #we need to get rid of any sentence that mentions family member, i.e 'father died of heart attack', otherwise we get false positives.
  #list all known family member types
  familyMembers = ["FATHER", "MOTHER", " SON ", " SONS ", "DAUGHTER", "BROTHER", "SISTER", "SISTERS", "GRANDFATHER", "GRANDFATHERS", "GRANDMOTHER", "GRANDMOTHERS", "GRANDPARENT", "CHILDREN", "UNCLE", "UNCLES", "AUNT", "AUNTS", "COUSIN", "FAMILY HISTORY OF", "SIBLING"]
  #check if any sentence contains any of the family member types
  for sentence in listedText:
    if any(member in sentence.upper() for member in familyMembers):
      #if found remove the sentence, otherwise it would be listed as if the patient had the condition, symptom, procedure etc.
      txtStripped = txtStripped.replace(sentence, "")
  return txtStripped

@anvil.server.callable
def AWS_MedicalComprehend(strMedText):
  # Define the local variables
  strMedMeasure = ""
  strMedFrequency = ""
  blnFalsePositive = 0
  #Connect to the AWS account and initiate the comprehend medical object
  client = boto3.client(service_name='comprehendmedical',
                        aws_access_key_id='AKIZVA6C6CDADPWEWQU5R72K87OI',
                        aws_secret_access_key='6E+xPCpwNWryF42uGyf84jjEB65o+HV1Sc3eMOJzaZlUJFxG',
                        region_name='ap-southeast-2')
  #clean up the medical text
  txtStripped = strip_Text(strMedText)

  #Pass the medical text to SNOMEDCT module to get SNOMED Codes and Concepts
  responseSNOMEDCT = client.infer_snomedct(Text = txtStripped)

  #Get the response which is a JSON file and store it in a dictionary
  dictSNOMEDCT = responseSNOMEDCT['Entities'];

  #Iterate the dictionary (the JSON file)
  for index, dictItem in enumerate(dictSNOMEDCT):
      #If the accuracy(Score) is less than 55% then ignore the entity, don't process it
    if dictItem['Score'] > 0.55:
    #Check if the entity has 'Traits' to see if the entity is a false positive or not
      if len(dictItem['Traits']) == 0:
        #if there is no 'Traits' entity then process the sentence and extract medical condition or procedure.
        populate_ConditionsProceduresDict(dictItem)
      else:
        for trait in dictItem['Traits']:
          # If 'trait' is a Negation and likelyhood of it is greater than 55% then it is a false positive.
          if trait['Name'] == 'NEGATION' and trait['Score'] > 0.55:
            blnFalsePositive = 1
        #If the entity is not 'false positive' then process the sentence and extract medical condition or procedure.
        if blnFalsePositive == 0:
          populate_ConditionsProceduresDict(dictItem)
        blnFalsePositive = 0

  #MEDICATION
  #Since the SNOMEDCT module doesn't extract medication data we need to call the core module
  #Pass the medical text to the core module
  result = client.detect_entities(Text = txtStripped)

  #Get the response which is a JSON file and store it in a dictionary
  dictEntities = result['Entities'];
  #Iterate the dictionary (the JSON file)
  for index, dictItem in enumerate(dictEntities):
    #If the accuracy(Score) is greater than 55% then process the entity and extract the medication data
    if dictItem['Score'] > 0.55:
      populate_MedicationDict(dictItem)
  
  #Store the result dictionaries in a list to return them to the UI
  #lstReturns = [lstCondition, lstProcedure,lstMedication, dictEntities, dictMedicalCondition, dictMedicalProcedure, dictMedicationGrid] #dictEntities/dictSNOMEDCT
  lstReturns = [ dictMedicalCondition, dictMedicalProcedure, dictMedicationGrid] 

  #return the dictionaries
  return (lstReturns)

def populate_ConditionsProceduresDict(dictItem):
  #This function populates the Conditions and Procedures dictionaries 
  global blnLastDictItemProcedure
  strOrgan = ''
  strDirection = ''
  strConditionType = ''
  intScore=0.00
  condition = {}
  procedure = {}

  #check if the entity is a 'Medical Condition'
  if dictItem['Category'] == "MEDICAL_CONDITION":
    strMedicalCondition = dictItem['Text']
    if 'Traits' in dictItem:
      for dictTrait in dictItem['Traits']:
        strConditionType = dictTrait['Name']
    if 'Attributes' in dictItem:
      dictMedConAttributes = dictItem['Attributes']
      for dictMedConAttribute in dictMedConAttributes:
        if dictMedConAttribute['RelationshipType'] == 'SYSTEM_ORGAN_SITE' and dictMedConAttribute['RelationshipScore'] > 0.5:
          if strOrgan == '':
            strOrgan = dictMedConAttribute['Text']
          else:
            strOrgan = strOrgan + ', ' + dictMedConAttribute['Text']
        if dictMedConAttribute['RelationshipType'] == 'DIRECTION':
          strDirection =  dictMedConAttribute['Text'] + ' '
    if strMedicalCondition.capitalize() not in lstCondition:
      if strOrgan == '':
        strMedicalCondition = strDirection.capitalize() + strMedicalCondition.capitalize() #Right Murmur instead of right murmur
      else:
        strMedicalCondition = strDirection.capitalize() + strMedicalCondition.capitalize() + ' (' + strOrgan + ')'
      lstCondition.append(strMedicalCondition)
    condition["Name"] = strMedicalCondition
    condition["Type"] = strConditionType
    if 'SNOMEDCTConcepts' in dictItem:
      dictSNOMEDCTConcepts = dictItem['SNOMEDCTConcepts']
      #for dictSNOMEDCTConcept in dictSNOMEDCTConcepts:
      condition["Code"] = dictSNOMEDCTConcepts[0]["Code"]
      condition["Description"] = dictSNOMEDCTConcepts[0]["Description"]
      #break
    else:      
      condition["Code"] = ''
      condition["Concept"] = ''
    dictMedicalCondition["condition"].append(condition)
  elif dictItem['Category'] == "TEST_TREATMENT_PROCEDURE":
    strTestName = dictItem['Text']
    strTestMeasure = ""
    if 'Attributes' in dictItem:
      dictTestAttributes = dictItem['Attributes']
      for dictTestAttribute in dictTestAttributes:
        if dictTestAttribute['RelationshipType'] == 'TEST_VALUE' and dictTestAttribute['RelationshipScore'] > 0.6:
          if strTestMeasure == '':
            strTestMeasure = dictTestAttribute['Text']
          else:
            strTestMeasure = strTestMeasure + '/' + dictTestAttribute['Text']
    if strTestMeasure == '':
      lstProcedure.append(strTestName.capitalize())
      procedure['Name'] = strTestName.capitalize()
      procedure['Measurement'] = ''
    else:
      lstProcedure.append(strTestName.capitalize() + ': ' + strTestMeasure)
      procedure['Name'] = strTestName.capitalize()
      procedure['Measurement'] = strTestMeasure

    if 'SNOMEDCTConcepts' in dictItem:
      dictSNOMEDCTConcepts = dictItem['SNOMEDCTConcepts']
      #for dictSNOMEDCTConcept in dictSNOMEDCTConcepts:
      procedure["Code"] = dictSNOMEDCTConcepts[0]["Code"]
      procedure["Description"] = dictSNOMEDCTConcepts[0]["Description"]
      #break
    else:
      procedure["Code"] = ''
      procedure["Concept"] = ''
    dictMedicalProcedure["procedure"].append(procedure)

    blnLastDictItemProcedure = 1

def populate_MedicationDict(dict_item):
  #This function populates the Medications dictionaries 
  dictMedication = {}
  if dict_item['Category'] == "MEDICATION":
    strMedName = dict_item['Text']
    dictMedication["Name"] = strMedName
    strMedMeasure = ""
    strMedFrequency = ""
    if 'Attributes' in dict_item:
      dictMedAttributes = dict_item['Attributes']
      for dictMedAttribute in dictMedAttributes:
        if dictMedAttribute['RelationshipType'] == 'DOSAGE':
          strMedMeasure = dictMedAttribute['Text']
          dictMedication["Dosage"] = strMedMeasure
        elif dictMedAttribute['RelationshipType'] == 'FREQUENCY':
          strMedFrequency = dictMedAttribute['Text']
          dictMedication["Frequency"] = strMedFrequency
    lstMedication.append(strMedName + ' ' + strMedMeasure + ' ' + strMedFrequency)

    dictMedicationGrid["medication"].append(dictMedication)


@anvil.server.callable
def create_ConditionsFile(lstConditionDicts):
  #This function creates the Conditions CSV file
  columns = ['Name', 'Type', 'Code', 'Description']
  with open("/tmp/Conditions_Data.csv", 'w',newline='') as csvfile:
    writer = csv.DictWriter(csvfile, fieldnames=columns)
    writer.writeheader()
    for key in lstConditionDicts:
      writer.writerow(key)
  X_ConditionsMedia = anvil.media.from_file('/tmp/Conditions_Data.csv', 'csv')
  #X_media = anvil.media.from_file('/tmp/countries.csv', 'csv')
  return X_ConditionsMedia

@anvil.server.callable
def create_MedicationsFile(lstMedicationsDicts):
  #This function creates the Medications CSV file
  columns = ['Name', 'Dosage', 'Frequency']
  with open("/tmp/Medications_Data.csv", 'w',newline='') as csvfile:
    writer = csv.DictWriter(csvfile, fieldnames=columns)
    writer.writeheader()
    for key in lstMedicationsDicts:
      writer.writerow(key)
  X_MedicationsMedia = anvil.media.from_file('/tmp/Medications_Data.csv', 'csv')
  #X_media = anvil.media.from_file('/tmp/countries.csv', 'csv')
  return X_MedicationsMedia

@anvil.server.callable
def create_ProceduresFile(lstProceduresDicts):
  #This function creates the Procedures CSV file
  columns = ['Name', 'Measurement', 'Code', 'Description']
  with open("/tmp/Procedures_Data.csv", 'w',newline='') as csvfile:
    writer = csv.DictWriter(csvfile, fieldnames=columns)
    writer.writeheader()
    for key in lstProceduresDicts:
      writer.writerow(key)
  X_ProceduresMedia = anvil.media.from_file('/tmp/Procedures_Data.csv', 'csv')
  #X_media = anvil.media.from_file('/tmp/countries.csv', 'csv')
  return X_ProceduresMedia

@anvil.server.callable
def Get_S3Files():
  #This function returns a list of files in an S3 bucket
  lstFiles = []
  session = boto3.Session(
                        aws_access_key_id='AFKIAWE6C6CDADPRTU5R7ZA2KOI',
                        aws_secret_access_key='6E+xPCpNWryFs42uwsGy8jd5djEB6+HV1ScgiMOJzaZU43JFxG')
  s3 = session.resource('s3')
  my_bucket = s3.Bucket('medletters')
  for my_bucket_object in my_bucket.objects.all():
    lstFiles.append(my_bucket_object.key)

  return lstFiles

@anvil.server.callable
def Process_S3File(strFile):
  #This function gets the file in the S3 bucket and downloads it to server for it to be displayed in the UI as the medical text.
  s3Client = boto3.client(service_name='s3',
                        aws_access_key_id='AFKIAWE6C6CDADPRTU5R7ZA2KOI',
                        aws_secret_access_key='6E+xPCpNWdyF42uwsGy8jd5djEB6+HV1ScgiMOJzaZU43JFxG',
                        region_name='ap-southeast-2')
  session = boto3.Session(
                        aws_access_key_id='AFKIAWE6C6CDADPRTU5R7ZA2KOI',
                        aws_secret_access_key='6E+xPCpNWryF42huwsGy8jd5djEB6+HV1ScgiMOJzaZU43JFxG')
  s3 = session.resource('s3')
  my_bucket = s3.Bucket('medletters')
  with open('/tmp/Letter.docx', 'wb') as f:
    s3Client.download_fileobj('medletters', strFile, f)
    strText = read_s3docx_file('/tmp/Letter.docx')
  return strText
  
@anvil.server.callable
def read_docx_file(objFile):
  #This function reads the content of a downloaded docx file
  #and returns the text to the UI 
  with open("/tmp/my-uploaded-file.docx", "wb") as f:
    f.write(objFile.get_bytes())

  doc = docx.Document("/tmp/my-uploaded-file.docx")
  fullText = []
  for para in doc.paragraphs:
    fullText.append(para.text)
  return '\n'.join(fullText)

@anvil.server.callable
def read_s3docx_file(objFile):
  #This function reads the content of a docx file in S3
  #and returns the text to the UI 
  doc = docx.Document(objFile)
  fullText = []
  for para in doc.paragraphs:
    fullText.append(para.text)
  return '\n'.join(fullText)


@anvil.server.callable
def read_pdf_file(objFile):
  #This function reads the content of a downloaded pdf file
  #and returns the text to the UI 
  with open("/tmp/my-uploaded-file.pdf", "wb") as f:
    f.write(objFile.get_bytes())

  with fitz.open("/tmp/my-uploaded-file.pdf") as doc:
    pdf_text = ""
    for page in doc:
      pdf_text += page.get_text()
    return pdf_text

